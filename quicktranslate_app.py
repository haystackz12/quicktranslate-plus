# quicktranslate_app.py
"""
QuickTranslate+  –  a feature-rich, one-file Streamlit translator
────────────────────────────────────────────────────────────────────
✔ Token / cost estimator                ✔ Batch upload + ZIP output
✔ Auto-detect source language           ✔ Progress bar per chunk
✔ Preserves original file type (txt / md / docx) on download
✔ Smart chunking for large texts        ✔ Optional privacy-mode (no caching)
✔ UI language toggle (English / Español)
✔ Adds one-paragraph summary of translation
────────────────────────────────────────────────────────────────────
Run locally:
    python -m venv venv && source venv/bin/activate
    pip install -r requirements.txt
    export OPENAI_API_KEY="sk-•••"
    streamlit run quicktranslate_app.py
"""

import os, io, zipfile, tempfile, textwrap, math
from typing import List, Tuple

import streamlit as st
from openai import OpenAI
from langdetect import detect
from docx import Document
from PyPDF2 import PdfReader
import tiktoken


# ╭──────────────────────────────────────────────────────────╮
# │ 1.  CONFIGURATION                                       │
# ╰──────────────────────────────────────────────────────────╯
client = OpenAI()                            # reads OPENAI_API_KEY env var
MODEL              = "gpt-3.5-turbo"
TOKENS_PER_CHUNK   = 3500                    # keep < 4k context
CHAR_LIMIT         = 30_000                  # quick portfolio limit
PRICE_PER_1K_TOKENS = 0.002                  # USD for gpt-3.5-turbo (June-2025)

ENCODING = tiktoken.encoding_for_model(MODEL)


# ╭──────────────────────────────────────────────────────────╮
# │ 2.  UI LANGUAGE STRINGS (en / es)                        │
# ╰──────────────────────────────────────────────────────────╯
UI_STR = {
    "en": {
        "title":        "🌍 QuickTranslate+",
        "uploader":     "Upload files (.txt, .md, .docx, .pdf)",
        "target_label": "Translate to",
        "privacy":      "Privacy mode (don’t cache content)",
        "translate":    "Translate",
        "cost_note":    "≈ {tokens} tokens · est. cost ${cost:.4f}",
        "char_warn":    "File exceeds the {limit} character limit.",
        "extract_err":  "Could not read text from {name}.",
        "summary_hd":   "🔎 Summary of translation",
        "download_all": "Download all translations (.zip)",
        "download_one": "Download translation",
        "done":         "✅ Completed!",
        "progress":     "Translating chunk {done}/{total}…",
    },
    "es": {
        "title":        "🌍 QuickTranslate+",
        "uploader":     "Sube archivos (.txt, .md, .docx, .pdf)",
        "target_label": "Traducir a",
        "privacy":      "Modo privacidad (no almacenar caché)",
        "translate":    "Traducir",
        "cost_note":    "≈ {tokens} tokens · costo aprox. ${cost:.4f}",
        "char_warn":    "El archivo supera el límite de {limit} caracteres.",
        "extract_err":  "No se pudo leer texto de {name}.",
        "summary_hd":   "🔎 Resumen de la traducción",
        "download_all": "Descargar todas las traducciones (.zip)",
        "download_one": "Descargar traducción",
        "done":         "✅ ¡Completado!",
        "progress":     "Traduciendo bloque {done}/{total}…",
    },
}


# ╭──────────────────────────────────────────────────────────╮
# │ 3.  HELPER FUNCTIONS                                     │
# ╰──────────────────────────────────────────────────────────╯
def extract_text(file) -> str:
    """Return plain text from uploaded file object."""
    name = file.name.lower()
    if name.endswith((".txt", ".md")):
        return file.read().decode("utf-8", errors="ignore")

    if name.endswith(".docx"):
        file.seek(0)
        doc = Document(file)
        return "\n".join(p.text for p in doc.paragraphs)

    if name.endswith(".pdf"):
        file.seek(0)
        reader = PdfReader(file)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    return ""


def token_count(text: str) -> int:
    return len(ENCODING.encode(text))


def split_into_chunks(text: str, max_tokens: int) -> List[str]:
    """Naive splitter by paragraphs, respecting max_tokens per chunk."""
    paragraphs = text.splitlines()
    chunks, current = [], []
    curr_tokens = 0

    for para in paragraphs:
        tks = token_count(para) + 1
        if curr_tokens + tks > max_tokens and current:
            chunks.append("\n".join(current))
            current, curr_tokens = [], 0
        current.append(para)
        curr_tokens += tks
    if current:
        chunks.append("\n".join(current))
    return chunks


def translate_chunk(chunk: str, target_lang: str) -> str:
    """Translate a single chunk synchronously."""
    system = (
        f"You are a professional translator. Translate the user's text into {target_lang}. "
        "Preserve paragraphs and line breaks. Output ONLY the translated text."
    )
    resp = client.chat.completions.create(
        model=MODEL,
        messages=[{"role": "system", "content": system},
                  {"role": "user",   "content": chunk}],
        temperature=0.3,
    )
    return resp.choices[0].message.content.strip()


def summarize_text(text: str, target_lang: str) -> str:
    """Return a short summary (3-4 sentences) of translated text."""
    prompt = (
        f"Provide a concise 3-4 sentence summary in {target_lang} of the following text:\n\n{text[:4000]}"
    )
    resp = client.chat.completions.create(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
    )
    return resp.choices[0].message.content.strip()


def build_output_bytes(translated: str, original_name: str) -> Tuple[bytes, str]:
    """Return (bytes, filename) preserving original extension when easy."""
    base, ext = os.path.splitext(original_name)
    if ext.lower() in (".txt", ".md"):
        return translated.encode("utf-8"), f"{base}_translated{ext}"

    if ext.lower() == ".docx":
        doc = Document()
        for line in translated.splitlines():
            doc.add_paragraph(line)
        tmp = io.BytesIO()
        doc.save(tmp)
        return tmp.getvalue(), f"{base}_translated.docx"

    # default to .txt
    return translated.encode("utf-8"), f"{base}_translated.txt"


# ╭──────────────────────────────────────────────────────────╮
# │ 4.  STREAMLIT INTERFACE                                  │
# ╰──────────────────────────────────────────────────────────╯
st.set_page_config(page_title="QuickTranslate+", page_icon="🌍", layout="centered")

# ── UI language toggle
ui_lang = st.sidebar.selectbox("🌐 UI Language", ["English", "Español"])
ui_key  = "es" if ui_lang.startswith("Esp") else "en"
T = UI_STR[ui_key]      # shortcut dict

st.title(T["title"])

# ── inputs
uploaded_files = st.file_uploader(
    T["uploader"], type=["txt", "md", "docx", "pdf"], accept_multiple_files=True
)
target_language = st.selectbox(
    T["target_label"],
    ["Spanish", "French", "German", "Chinese", "Japanese", "Portuguese",
     "Italian", "Korean", "Russian", "Arabic", "Hindi"],
)
privacy_mode = st.checkbox(T["privacy"], value=False)

if st.button(T["translate"]) and uploaded_files:
    all_outputs = []

    for file in uploaded_files:
        # ---- extraction
        text = extract_text(file)
        if not text.strip():
            st.error(T["extract_err"].format(name=file.name))
            continue
        if len(text) > CHAR_LIMIT:
            st.warning(T["char_warn"].format(limit=CHAR_LIMIT))
            continue

        # ---- diagnostics: token + cost
        total_tokens = token_count(text)
        est_cost     = total_tokens / 1000 * PRICE_PER_1K_TOKENS
        st.info(T["cost_note"].format(tokens=total_tokens, cost=est_cost))

        # ---- source language detection
        try:
            src_lang = detect(text[:2000])
        except Exception:
            src_lang = "unknown"

        # ---- chunking
        chunks = split_into_chunks(text, TOKENS_PER_CHUNK)

        # ---- translate with progress
        prog = st.progress(0, text=T["progress"].format(done=0, total=len(chunks)))
        translated_parts = []
        for i, chunk in enumerate(chunks, 1):
            translated_parts.append(translate_chunk(chunk, target_language))
            prog.progress(i / len(chunks),
                          text=T["progress"].format(done=i, total=len(chunks)))
        prog.empty()

        translated_full = "\n".join(translated_parts)

        # ---- summary
        summary = summarize_text(translated_full, target_language)
        st.subheader(T["summary_hd"])
        st.write(summary)

        # ---- build output
        bytes_out, out_name = build_output_bytes(translated_full, file.name)
        st.download_button(
            T["download_one"], data=bytes_out, file_name=out_name, mime="application/octet-stream"
        )

        all_outputs.append((out_name, bytes_out))

    # ---- zip batch
    if len(all_outputs) > 1:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".zip") as tmpzip:
            with zipfile.ZipFile(tmpzip, "w") as z:
                for fname, b in all_outputs:
                    z.writestr(fname, b)
        with open(tmpzip.name, "rb") as fzip:
            zbytes = fzip.read()
        st.download_button(
            T["download_all"],
            data=zbytes,
            file_name="translations.zip",
            mime="application/zip",
        )

    st.success(T["done"])
